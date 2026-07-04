from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "数字周参-近15天文章检索处理报告.docx"


def add_hyperlink(paragraph, text, url):
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for margin in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
    setattr(section, margin, Inches(1))

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for style_name, size, color in [
    ("Heading 1", 16, "2E74B5"),
    ("Heading 2", 13, "2E74B5"),
]:
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)

title = doc.add_paragraph()
r = title.add_run("数字周参近15天文章检索处理报告")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor.from_string("0B2545")

p = doc.add_paragraph()
p.add_run("检索时间窗：").bold = True
p.add_run("2026-06-09 至 2026-06-23（含，按当前日期 2026-06-23 计算最近15天）")

p = doc.add_paragraph()
p.add_run("处理结论：").bold = True
p.add_run(
    "在当前可核验的公开来源、搜索结果和已保存公众号转载链路中，未发现“数字周参”在上述15天窗口内发布的新文章。"
    "因此本次不生成具体文章的详细精读版，改为形成检索处理报告。"
)

doc.add_heading("一、检索范围与方法", level=1)
for item in [
    "按项目 AGENTS.md 中的微信公众号处理流程执行，先核验时间顺序与来源，再决定是否生成精读 Word。",
    "检索关键词包括“数字周参”“IT老周”“2026年6月”“mp.weixin.qq.com”等组合。",
    "交叉查看公开转载页、搜索引擎结果、此前已核验的数字周参文章链接。",
    "因微信公众号原站存在环境验证和反爬限制，不能把无法打开的页面当作有效发布时间依据。",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("二、检索结果", level=1)
table = doc.add_table(rows=1, cols=4)
headers = ["项目", "结果", "时间", "说明"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    shade_cell(table.rows[0].cells[i], "F2F4F7")

rows = [
    (
        "最近15天内文章",
        "未发现可核验文章",
        "2026-06-09 至 2026-06-23",
        "公开检索未命中数字周参在此窗口内的新文章。",
    ),
    (
        "当前可核验最近文章",
        "教育信息化行业，不只是没钱了",
        "2026-05-19 23:12",
        "已由公开转载页记录公众号、作者、发布时间和原文链接。",
    ),
    (
        "更早一篇可核验文章",
        "当信息中心变成“兜底部门”，组织已经出问题了",
        "2026-03-17",
        "公开转载页标注来源于数字周参，作者 IT老周。",
    ),
]
for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

doc.add_heading("三、可核验来源", level=1)
p = doc.add_paragraph()
p.add_run("最近可核验文章：").bold = True
add_hyperlink(
    p,
    "教育信息化行业，不只是没钱了（公开转载页）",
    "https://gitcode.csdn.net/6a0d1c73662f9a54cb75cb43.html",
)

p = doc.add_paragraph()
p.add_run("原文链接：").bold = True
add_hyperlink(
    p,
    "https://mp.weixin.qq.com/s/1MV3XHG2ne77sQT6hz-l2Q",
    "https://mp.weixin.qq.com/s/1MV3XHG2ne77sQT6hz-l2Q",
)

p = doc.add_paragraph()
p.add_run("更早文章核验页：").bold = True
add_hyperlink(
    p,
    "当信息中心变成“兜底部门”，组织已经出问题了",
    "https://wechatview.bluecatbot.com/%E8%87%AA%E4%B8%BB%E5%8F%AF%E6%8E%A7%E6%96%B0%E9%B2%9C%E4%BA%8B/2026-03-17-22-00%E5%BD%93%E4%BF%A1%E6%81%AF%E4%B8%AD%E5%BF%83%E5%8F%98%E6%88%90%E2%80%9C%E5%85%9C%E5%BA%95%E9%83%A8%E9%97%A8%E2%80%9D%EF%BC%8C%E7%BB%84%E7%BB%87%E5%B7%B2%E7%BB%8F%E5%87%BA%E9%97%AE%E9%A2%98%E4%BA%86/index.html",
)

doc.add_heading("四、后续建议", level=1)
for item in [
    "如果用户确认微信客户端内能看到最近15天内数字周参有新文章，可提供文章链接或截图标题，我再按详细精读版流程处理。",
    "若需要持续跟踪，可建立一个固定检索窗口：每次先确认最近文章列表，再只处理新增文章。",
    "本次不上传 IMA 知识库，因为当前请求未指定入库，且没有发现新的文章精读文档可入库。",
]:
    doc.add_paragraph(item, style="List Number")

doc.save(OUT)
print(OUT)
