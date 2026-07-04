from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "数字周参-教育信息化行业详细精读版.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url):
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullet(title, body):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title + "：")
    r.bold = True
    p.add_run(body)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for margin in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
    setattr(section, margin, Inches(1))

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
r = title.add_run("教育信息化行业，不只是没钱了")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor.from_string("0B2545")

sub = doc.add_paragraph()
sub.add_run("数字周参文章详细精读版").bold = True
sub.add_run(" | 原文作者：IT老周 | 公众号：数字周参 | 原文时间：2026年5月19日 23:12")

p = doc.add_paragraph()
p.add_run("原文链接：").bold = True
add_hyperlink(p, "https://mp.weixin.qq.com/s/1MV3XHG2ne77sQT6hz-l2Q", "https://mp.weixin.qq.com/s/1MV3XHG2ne77sQT6hz-l2Q")

note = doc.add_paragraph()
note.add_run("说明：").bold = True
note.add_run(
    "本文档是对原文观点、结构和论证路径的详细精读与改写，覆盖原文主要信息，但不逐字复制全文。"
    "适合用于内部学习、行业判断、销售/产品/转型讨论。"
)

doc.add_heading("一、全文主旨", level=1)
add_para(
    "文章的核心判断是：教育信息化行业并不是没有机会了，而是粗放增长的时代过去了。"
    "过去企业依靠政策窗口、专项资金、区域批量建设和学校集中采购，往往能获得较快增长；"
    "但现在预算更紧、监管更严、采购更谨慎、客户更务实，行业进入了从增量建设到存量经营的阶段。"
)
add_para(
    "因此，从业者的痛感不是单纯来自“没钱”，而是来自旧打法失效。"
    "学校仍然有信息化、数字化、智能化需求，但这些需求不再自动变成大项目，也不再天然支持宏大平台。"
    "企业必须证明自己能在真实场景里解决具体问题。"
)

doc.add_heading("二、文章开篇的问题意识", level=1)
add_para(
    "作者先从读者私信切入：一线销售关心行业是否还值得做，大厂和科技公司关心现在进入是否太晚，"
    "传统装备、功能室、实验室和集成公司关心是否要转型，求职者则关心行业岗位和职业前途。"
)
add_para(
    "这些问题表面不同，背后都指向同一个焦虑：教育信息化是不是已经不行了。"
    "作者的回答不是简单唱衰，也不是鼓励盲目乐观，而是把行业变化拆开看。"
    "机会仍在，只是机会的形态、获取方式和能力要求都变了。"
)

doc.add_heading("三、从增量时代到存量时代", level=1)
add_para(
    "过去的教育信息化更像建设型行业。政策文件、专项资金、区域工程、学校改造、平台建设等因素叠加，"
    "让很多企业能围绕“建系统、上平台、配设备、做示范”获得项目。"
)
add_para(
    "现在的行业则更像经营型行业。客户不再只问有没有系统，而是问系统能不能真正被老师、学生、校长、处室和教育局用起来；"
    "能不能减少工作负担、提升管理效率、降低风险、支撑考核；能不能在预算受限时仍然体现价值。"
)
add_para(
    "这也是作者反复强调的变化：以前是增量时代，现在是存量时代；以前是建设逻辑，现在是经营逻辑；"
    "以前是大水漫灌，现在是小口径滴灌。企业不能再只等政策红利，而要进入学校真实业务。"
)

doc.add_heading("四、教育信息化不是完全市场化行业", level=1)
add_para(
    "文章提醒，教育信息化不能简单当作普通 To B 生意。学校看起来有需求，并不代表马上能采购；"
    "校长认可、信息中心认可、处室觉得系统不好用，也不等于预算已经存在。"
)
add_para(
    "教育系统采购受到多重因素制约：政策方向、财政预算、专项资金、地方支付能力、项目立项流程、审计监管、组织权责和一把手推动意愿。"
    "任何一环卡住，需求就可能停留在口头层面。"
)
add_para(
    "文章里最值得记住的一点是：学校有需求，不代表有项目；有项目，不代表有预算；有预算，不代表能马上花；"
    "能花，也不代表花在你的方向；流程走完，也不代表回款快。"
)

doc.add_heading("五、为什么很多企业觉得难", level=1)
add_bullet(
    "预算环境变化",
    "财政支出更谨慎，教育信息化项目不再天然被优先支持。很多地方仍有需求，但项目节奏变慢、金额变小、审批更细。"
)
add_bullet(
    "客户决策更务实",
    "学校和教育局更关心系统是否真能解决问题，而不是方案名字是否宏大、展示页面是否漂亮。"
)
add_bullet(
    "监管和审计更严格",
    "过去一些依赖关系、包装、概念和渠道的做法空间缩小，项目合规性、必要性、验收结果和资金使用都会被更认真地审视。"
)
add_bullet(
    "旧产品失去吸引力",
    "泛平台、泛门户、泛大屏、泛数据中台如果不能落到日常事务，很容易被客户认为“有也行，没有也行”。"
)
add_bullet(
    "交付和运营压力上升",
    "客户不只要上线，还要用起来、见成效、有人负责维护。企业必须承担更多长期服务责任。"
)

doc.add_heading("六、平台型方案的退潮", level=1)
add_para(
    "文章对“大而全”的智慧校园、区域平台、综合解决方案保持警惕。过去这类方案容易成为项目包装入口，"
    "因为它们听起来完整、宏大，也便于做汇报和展示。"
)
add_para(
    "但在预算收紧和客户更理性的情况下，平台本身不再构成充分理由。客户会进一步追问：谁每天使用？解决哪个流程？"
    "减少了什么工作？沉淀了什么数据？能不能验收？后续谁维护？如果回答不上来，大平台就会变成空概念。"
)
add_para(
    "这不是说平台没有价值，而是平台必须从具体场景长出来。先有高频事务、明确责任、稳定流程和可衡量结果，"
    "再谈数据汇聚、平台能力和区域扩展，才更符合现在的行业节奏。"
)

doc.add_heading("七、真正的机会在哪里", level=1)
add_para(
    "作者认为机会仍在，但机会更小、更细、更贴近现场。客户愿意为具体问题买单，而不是为抽象概念买单。"
)
for title_text, body in [
    ("学生安全管理", "围绕出入校、宿舍、请假、异常预警、家校沟通、责任追踪等高频风险点形成闭环。"),
    ("考试与学情分析", "不只是报表展示，而是帮助老师、年级组、教务处看清问题、安排教学、追踪改进。"),
    ("设备与资产运维", "解决设备多、维护难、责任不清、报修低效、账实不一致等问题。"),
    ("实训与职教管理", "职业院校对实训室、课程、设备、耗材、考核和安全有真实管理需求。"),
    ("教学与行政流程", "把请假、审批、排课、督导、评价、材料归档等事务做得更顺。"),
]:
    add_bullet(title_text, body)

doc.add_heading("八、三类企业会更难", level=1)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table)
for i, h in enumerate(["企业类型", "主要问题", "应对方向"]):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, "F2F4F7")
    set_cell_margins(cell)
rows = [
    ("纯商贸型", "只做倒货、拼价格、拼渠道，缺少产品沉淀和服务能力。", "补齐场景方案、交付实施、运维服务和客户经营能力。"),
    ("纯关系型", "过度依赖熟人和资源，面对预算收紧、监管加强时抗风险能力弱。", "从关系驱动转向价值驱动，用可验证成果增强客户信任。"),
    ("纯概念型", "擅长包装新名词，但产品难落地，客户很难看到实际效果。", "减少空泛叙事，把概念拆成可使用、可验收、可复购的小产品。"),
]
for row_data in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row_data):
        cells[i].text = text
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.add_heading("九、行业能力要求的变化", level=1)
add_para(
    "过去企业可能只需要会拿项目、会讲方案、会交付系统。现在则需要更复合的能力：既懂教育政策和财政采购，"
    "又懂学校内部组织和日常流程；既能做产品，又能做运营；既能卖出去，也能让客户持续用。"
)
add_bullet(
    "销售能力",
    "销售不能只问项目预算和采购时间，还要会判断客户真实痛点、组织结构、内部推动人、预算来源、验收标准和回款风险。"
)
add_bullet(
    "方案能力",
    "方案不能只堆功能模块，要围绕具体业务路径，把用户、场景、流程、数据、责任和结果说清楚。"
)
add_bullet(
    "产品能力",
    "产品不应追求“大而全”，而应优先打穿高频刚需场景，形成可重复交付的小闭环。"
)
add_bullet(
    "交付能力",
    "交付不只是安装部署，还包括培训、上线、使用跟踪、问题响应和效果复盘。"
)
add_bullet(
    "经营能力",
    "企业要从一次性项目收入，逐步转向持续服务、运维、轻量订阅、场景扩展和客户复购。"
)

doc.add_heading("十、对不同角色的详细建议", level=1)
for role, advice in [
    ("对一线销售", "少讲宏大概念，多做场景诊断。拜访时要问清楚学校目前哪个环节最痛、谁负责、为什么现在要解决、预算从哪里来、验收看什么。真正有效的销售，是帮客户把问题定义清楚。"),
    ("对产品经理", "不要把教育信息化理解成通用 SaaS 换皮。学校有自己的组织结构、权限边界、流程惯性和考核压力。产品设计要尊重现场，不要让老师为了系统额外劳动。"),
    ("对方案岗", "方案需要从“漂亮文档”变成“推进工具”。它既要能对上级汇报，也要能帮助客户内部达成共识，还要能指导交付落地。"),
    ("对传统集成商", "不能只靠硬件和渠道。未来更重要的是软硬结合、运维服务、数据沉淀和客户长期经营。"),
    ("对求职者", "行业不是没有岗位，而是岗位更挑人。优先选择有真实客户、有产品迭代、有长期服务能力的公司，并提升教育业务理解能力。"),
]:
    doc.add_heading(role, level=2)
    add_para(advice)

doc.add_heading("十一、企业转型清单", level=1)
for item in [
    "重新梳理客户需求，把“智慧校园”“教育数字化”等大概念拆成具体事务场景。",
    "筛选最有付费可能、最高频、最刚需、最容易验收的场景作为切入口。",
    "为每个场景定义清楚用户、流程、数据、责任人、结果指标和交付边界。",
    "减少只为演示而做的功能，把资源投入到能被客户每天使用的能力上。",
    "建立客户上线后的持续服务机制，跟踪使用率、问题响应和效果反馈。",
    "销售、方案、产品、交付共同复盘项目，不让一线经验停留在个人身上。",
    "谨慎评估大项目风险，尤其关注付款周期、验收标准和后续维护成本。",
]:
    doc.add_paragraph(item, style="List Number")

doc.add_heading("十二、精读结论", level=1)
add_para(
    "这篇文章真正想提醒从业者的是：不要用过去的成功经验解释今天的行业。"
    "教育信息化仍然存在真实需求，但行业已经从“有政策就能上项目”的阶段，转到“能解决具体事务才有价值”的阶段。"
)
add_para(
    "对企业而言，未来的竞争不只是资源竞争，也不是概念竞争，而是场景理解、产品打磨、交付运营和长期经营能力的竞争。"
    "越是预算紧、客户谨慎的时候，越能检验企业是不是真的懂教育现场。"
)
add_para(
    "所以，最重要的不是判断这个行业还能不能做，而是判断自己有没有能力换一种方式做："
    "从等风口，转向做现场；从卖项目，转向经营客户；从讲平台，转向解决具体事务。"
)

doc.add_heading("十三、来源", level=1)
p = doc.add_paragraph()
p.add_run("公众号原文：").bold = True
add_hyperlink(p, "教育信息化行业，不只是没钱了", "https://mp.weixin.qq.com/s/1MV3XHG2ne77sQT6hz-l2Q")
p = doc.add_paragraph()
p.add_run("公开转载核验页：").bold = True
add_hyperlink(p, "AtomGit/CSDN 转载页", "https://gitcode.csdn.net/6a0d1c73662f9a54cb75cb43.html")

doc.save(OUT)
print(OUT)
