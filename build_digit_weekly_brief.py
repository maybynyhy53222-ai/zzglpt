from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "数字周参-最新文章摘录与总结.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def style_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for margin in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
    setattr(section, margin, Inches(1))

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    st = styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
run = title.add_run("数字周参最新文章摘录与总结")
style_run(run, bold=True, color="0B2545", size=22)

subtitle = doc.add_paragraph()
subtitle.add_run("文章来源：").bold = True
subtitle.add_run("数字周参 | 作者：IT老周 | 发布时间：2026年5月19日 23:12 | 发布地：广东")

p = doc.add_paragraph()
p.add_run("原文链接：").bold = True
add_hyperlink(p, "https://mp.weixin.qq.com/s/1MV3XHG2ne77sQT6hz-l2Q", "https://mp.weixin.qq.com/s/1MV3XHG2ne77sQT6hz-l2Q")

p = doc.add_paragraph()
p.add_run("检索说明：").bold = True
p.add_run(
    "按用户要求尝试从已登录微信查找“数字周参”最新文章；同时核验到公开转载页记录该文为公众号“数字周参”的文章。"
    "本文档不整篇转载原文，仅保留必要短摘录和基于阅读的结构化总结。"
)

doc.add_heading("一、文章基本信息", level=1)
table = doc.add_table(rows=5, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table)
items = [
    ("标题", "教育信息化行业，不只是没钱了"),
    ("公众号", "数字周参"),
    ("作者", "IT老周"),
    ("时间", "2026年5月19日 23:12"),
    ("主题", "教育信息化行业从增量建设转向存量经营后的机会、风险和企业转型方向"),
]
for row, (k, v) in zip(table.rows, items):
    row.cells[0].text = k
    row.cells[1].text = v
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
    set_cell_shading(row.cells[0], "F2F4F7")

doc.add_heading("二、原文短摘录", level=1)
p = doc.add_paragraph()
p.add_run("合规说明：").bold = True
p.add_run("以下仅摘取少量核心句，便于识别文章主旨；未复制全文。")
for quote in [
    "教育信息化行业不是没机会了，而是大水漫灌的机会没了。",
    "以前是增量时代。现在是存量时代。",
]:
    q = doc.add_paragraph(style=None)
    q.paragraph_format.left_indent = Inches(0.25)
    q.paragraph_format.space_after = Pt(4)
    r = q.add_run(f"“{quote}”")
    r.italic = True
    r.font.color.rgb = RGBColor.from_string("555555")

doc.add_heading("三、核心结论", level=1)
lead = doc.add_paragraph()
lead.add_run("一句话概括：").bold = True
lead.add_run(
    "文章认为教育信息化没有消失，但行业红利形态已经改变：过去依赖政策窗口、专项资金和批量建设的增长方式退潮，"
    "未来更看重企业能否围绕学校真实事务场景做出可落地、可运营、可持续的小闭环产品。"
)

doc.add_heading("四、结构化总结", level=1)
summary_points = [
    (
        "行业不是简单“没钱”，而是底层逻辑变了",
        "文章把当前困难归因于行业从政策驱动的增量建设，转向预算收紧、存量优化和结果导向的经营阶段。客户仍有需求，但项目形成、预算释放、采购流程和回款节奏都更复杂。"
    ),
    (
        "教育信息化不是完全市场化生意",
        "学校和教育系统采购受到政策方向、财政预算、专项资金、地方支付能力、审批流程、审计监管和一把手意愿等多重因素影响。学校“想要”不等于能立项，也不等于能快速付款。"
    ),
    (
        "平台型、大而全方案的吸引力下降",
        "过去很多企业习惯讲智慧校园、区域平台、综合解决方案，但现在客户更谨慎，愿意为能解决具体事务的小场景买单。泛平台、泛概念、泛展示型项目更难推动。"
    ),
    (
        "机会转向真实场景里的小闭环",
        "文章重点强调学生安全管理、考试分析、设备运维、教学管理、实训管理等明确场景。能把流程、数据、责任人、结果反馈做完整的产品，更有生存空间。"
    ),
    (
        "三类企业会更难",
        "纯商贸型公司缺少产品与服务沉淀，纯关系型公司在监管和预算压力下空间收窄，纯概念型公司难以证明价值。行业会逼企业补足产品、交付、运营和服务能力。"
    ),
    (
        "转型方向是从“卖项目”到“陪客户把事做成”",
        "企业要理解学校的组织结构和业务流程，把产品嵌进日常工作，而不是停留在演示、宣传和一次性交付。销售、方案、产品、交付都需要更懂教育现场。"
    ),
]
for title_text, body in summary_points:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.add_run(title_text + "：").bold = True
    p.add_run(body)

doc.add_heading("五、对不同读者的启示", level=1)
audience_table = doc.add_table(rows=1, cols=3)
audience_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(audience_table)
headers = ["读者类型", "主要判断", "建议动作"]
for i, h in enumerate(headers):
    cell = audience_table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, "F2F4F7")
    set_cell_margins(cell)

audiences = [
    ("一线销售", "不能只盯项目名和预算口径，要理解学校真实事务、流程痛点和付款路径。", "从“介绍产品”转向“诊断场景”，把客户内部流程、负责人、预算来源和验收指标问清楚。"),
    ("产品/方案岗", "大而全方案不再天然有吸引力，产品价值必须压到具体任务里。", "围绕学生安全、考试分析、设备运维、实训管理等高频场景设计小闭环。"),
    ("集成商/装备商", "单纯卖硬件、拼关系、拼资质会越来越难。", "补交付、运维、数据和软件能力，形成长期服务，而不是一次性供货。"),
    ("求职者", "行业仍有岗位，但更需要复合能力。", "优先选择有真实产品、持续客户和场景深度的公司，提升教育业务理解和项目推进能力。"),
]
for row_data in audiences:
    cells = audience_table.add_row().cells
    for i, text in enumerate(row_data):
        cells[i].text = text
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.add_heading("六、可执行要点", level=1)
actions = [
    "评估一个教育信息化机会时，先问预算来源、立项路径、采购周期、验收标准和回款风险。",
    "把方案拆成可被学校日常部门使用的小场景，避免只做宏大叙事。",
    "产品设计要能沉淀流程数据和结果反馈，让客户看到管理效率或风险控制的变化。",
    "企业经营上要减少对单一政策窗口和一次性大项目的依赖，增强续费、运维和轻量服务能力。",
    "团队能力建设上，销售、产品、交付都需要懂教育组织、懂现场流程、懂财政采购逻辑。",
]
for item in actions:
    doc.add_paragraph(item, style="List Number")

doc.add_heading("七、我的简评", level=1)
doc.add_paragraph(
    "这篇文章的价值不在于判断教育信息化“好”或“不好”，而是把行业困境拆成了更可操作的变化："
    "预算更谨慎、采购更慢、概念型平台退潮、客户更看重真实事务结果。对从业者来说，关键不是退出还是进入，"
    "而是尽快从红利思维切换到经营思维。谁能更贴近学校现场，谁就更可能在存量阶段留下来。"
)

doc.add_heading("八、来源", level=1)
p = doc.add_paragraph()
p.add_run("公众号原文：").bold = True
add_hyperlink(p, "教育信息化行业，不只是没钱了", "https://mp.weixin.qq.com/s/1MV3XHG2ne77sQT6hz-l2Q")
p = doc.add_paragraph()
p.add_run("公开转载核验页：").bold = True
add_hyperlink(p, "AtomGit/CSDN 转载页", "https://gitcode.csdn.net/6a0d1c73662f9a54cb75cb43.html")

doc.save(OUT)
print(OUT)
